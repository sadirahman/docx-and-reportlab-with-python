from docx import Document
from docx.shared import Inches

document = Document()
document.add_heading("Introduction")
document.add_paragraph("this is introduction")
document.save('demee.docx')